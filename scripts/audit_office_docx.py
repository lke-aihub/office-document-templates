#!/usr/bin/env python3
"""Audit geometry, core typography, and unresolved placeholders in a DOCX."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from format_profiles import PROFILES


if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")


EMU_PER_MM = 36000


def mm(value) -> float:
    return round(value / EMU_PER_MM, 2)


def style_east_asia(style) -> str | None:
    rpr = style.element.rPr
    if rpr is None or rpr.rFonts is None:
        return None
    return rpr.rFonts.get(qn("w:eastAsia"))


def close(actual: float, expected: float, tolerance: float = 0.25) -> bool:
    return abs(actual - expected) <= tolerance


def audit(path: Path, profile_name: str) -> dict:
    profile = PROFILES[profile_name]
    doc = Document(path)
    section = doc.sections[0]
    expected_page = profile["page"]
    actual_page = {
        "width_mm": mm(section.page_width),
        "height_mm": mm(section.page_height),
        "top_mm": mm(section.top_margin),
        "bottom_mm": mm(section.bottom_margin),
        "left_mm": mm(section.left_margin),
        "right_mm": mm(section.right_margin),
    }
    issues: list[str] = []
    for key, actual in actual_page.items():
        expected = expected_page[key]
        if not close(actual, expected):
            issues.append(f"{key}: expected {expected} mm, got {actual} mm")

    for style_name, expected in (("Normal", profile["body"]), ("Title", profile["title"])):
        style = doc.styles[style_name]
        size = style.font.size.pt if style.font.size else None
        if size is None or abs(size - expected["size_pt"]) > 0.1:
            issues.append(f"{style_name} size: expected {expected['size_pt']} pt, got {size}")
        actual_font = style_east_asia(style)
        if actual_font != expected["font"]:
            issues.append(f"{style_name} font: expected {expected['font']}, got {actual_font}")

    all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    placeholders = sorted({part for part in all_text.split() if "[待补充" in part})
    if "[待补充" in all_text and not placeholders:
        placeholders = ["[待补充：……]"]

    return {
        "path": str(path.resolve()),
        "profile": profile_name,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "page": actual_page,
        "placeholders_present": "[待补充" in all_text,
        "placeholder_tokens": placeholders,
        "issues": issues,
        "passed": not issues,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--profile", required=True, choices=PROFILES)
    args = parser.parse_args()
    if not args.docx.is_file():
        raise SystemExit(f"File not found: {args.docx}")
    result = audit(args.docx, args.profile)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if result["passed"] else 1


if __name__ == "__main__":
    sys.exit(main())

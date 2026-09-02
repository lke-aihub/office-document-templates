#!/usr/bin/env python3
"""Generate deterministic Chinese office DOCX files from a compact JSON payload.

The calling agent writes the polished text once, shows that text in chat, and
passes the same content here for Word output. Run with the Python runtime
returned by Codex workspace dependencies so ``python-docx`` is available.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from format_profiles import FORMAT_SOURCES, PROFILES, TYPE_DEFAULT_PROFILE, resolve_profile


if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")


DOCUMENT_TYPES = tuple(TYPE_DEFAULT_PROFILE)

SCHEMAS = {
    "leave-note": {
        "title": "请假条",
        "recipient": "王经理",
        "leave_type": "年假",
        "reason": "处理个人事务",
        "start_date": "2026年8月26日上午",
        "end_date": "2026年8月27日下午",
        "duration": "2天",
        "handover": "已将客户跟进事项交接给李明",
        "contact": "13800000000",
        "signer": "张三",
        "department": "市场部",
        "date": "2026年8月25日",
    },
    "request": {
        "title": "关于申请购置会议设备的请示",
        "recipient": "公司领导",
        "background": "现有设备无法满足远程会议需要。",
        "proposal": "拟购置视频会议设备一套，预算为人民币[待补充：金额]元。",
        "request": "现申请批准上述采购方案。",
        "closing": "妥否，请批示。",
        "signer": "行政部",
        "date": "2026年8月25日",
        "official": False,
    },
    "notice": {
        "title": "关于召开项目推进会的通知",
        "recipient": "各项目组成员",
        "intro": "为推进重点项目按期交付，现将会议安排通知如下：",
        "sections": [
            {"heading": "会议时间", "paragraphs": ["2026年8月28日14:00"]},
            {"heading": "会议地点", "paragraphs": ["第一会议室"]},
            {"heading": "参会要求", "paragraphs": ["请于会前提交最新进度及风险清单。"]},
        ],
        "closing": "特此通知。",
        "signer": "项目管理办公室",
        "date": "2026年8月25日",
        "official": False,
    },
    "application": {
        "title": "培训经费申请书",
        "recipient": "人力资源部",
        "identity": "本人为产品部产品经理",
        "request": "申请参加[待补充：培训名称]",
        "reason": "该培训与当前项目能力建设直接相关",
        "details": "培训时间、费用及课程安排见附件。",
        "closing": "恳请审核批准。",
        "signer": "张三",
        "department": "产品部",
        "date": "2026年8月25日",
    },
    "formal-email": {
        "subject": "项目方案确认｜请于8月28日前反馈",
        "to": "王经理",
        "cc": "项目组",
        "greeting": "王经理，您好：",
        "paragraphs": ["现将项目方案发送给您审阅。", "本次更新主要涉及实施范围和交付时间。"],
        "action": "烦请于2026年8月28日前回复确认。",
        "closing": "感谢您的支持与配合。",
        "signer": "张三",
        "department": "产品部",
        "contact": "13800000000",
        "attachments": ["项目方案V2.0.docx"],
    },
    "travel-report": {
        "title": "关于赴上海开展客户调研的出差报告",
        "summary": "本次出差完成三家重点客户访谈，形成需求清单和后续推进计划。",
        "basics": {"时间": "2026年8月18日至20日", "地点": "上海", "人员": "张三、李四", "目标": "客户需求调研"},
        "itinerary": ["拜访A客户并访谈业务负责人。", "组织解决方案交流会。"],
        "results": ["确认三项核心需求。", "形成下一阶段联合验证计划。"],
        "issues": ["部分接口数据仍待客户确认。"],
        "suggestions": ["由技术团队于8月30日前完成接口清单复核。"],
        "follow_up": ["张三负责客户确认；李四负责技术评估。"],
        "signer": "张三",
        "department": "产品部",
        "date": "2026年8月25日",
    },
    "weekly-report": {
        "title": "产品部2026年第34周工作周报",
        "summary": "本周核心版本按计划上线，客户验证进度正常。",
        "achievements": ["完成V2.1版本发布。", "完成重点客户需求评审。"],
        "issues": ["测试环境资源存在排期冲突。"],
        "plans": ["推进V2.2需求拆解。", "完成客户验证问题闭环。"],
        "support": ["请协调测试环境在周三前释放。"],
        "signer": "产品部",
        "date": "2026年8月25日",
    },
    "monthly-report": {
        "title": "产品部2026年8月工作月报",
        "summary": "本月完成核心版本交付，整体进度符合计划。",
        "achievements": ["完成两次版本发布。"],
        "issues": ["客户验收周期较原计划延长。"],
        "plans": ["完成验收材料补充并推进签署。"],
        "support": [],
        "signer": "产品部",
        "date": "2026年8月25日",
    },
    "work-report": {
        "title": "重点项目阶段工作汇报",
        "summary": "项目整体可控，核心里程碑已完成，需关注测试资源风险。",
        "achievements": ["核心功能开发完成。"],
        "issues": ["测试资源不足可能影响验收。"],
        "plans": ["完成全量回归测试并准备验收。"],
        "support": ["请协调一组专项测试资源。"],
        "signer": "项目组",
        "date": "2026年8月25日",
    },
    "official-document": {
        "title": "关于进一步规范材料报送工作的通知",
        "recipient": "各有关单位",
        "intro": "为进一步提高材料报送质量，现将有关事项通知如下：",
        "sections": [
            {"heading": "一、明确报送内容", "paragraphs": ["按要求完整填写有关信息。"]},
            {"heading": "二、严格报送时限", "paragraphs": ["请于[待补充：日期]前完成报送。"]},
        ],
        "closing": "特此通知。",
        "signer": "[待补充：发文机关]",
        "date": "2026年8月25日",
        "official": True,
    },
}

ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

CN_NUMERALS = "一二三四五六七八九十"


def as_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        return [value] if value.strip() else []
    if isinstance(value, Iterable) and not isinstance(value, (dict, bytes)):
        return [str(item) for item in value if str(item).strip()]
    return [str(value)]


def require_or_placeholder(payload: dict, key: str, label: str) -> str:
    value = str(payload.get(key, "")).strip()
    return value or f"[待补充：{label}]"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, spec: dict, *, bold: bool | None = None, color: str | None = None) -> None:
    east_asia = spec.get("font", "宋体")
    ascii_font = spec.get("ascii_font", "Calibri")
    run.font.name = ascii_font
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run.font.size = Pt(spec.get("size_pt", 12))
    run.font.bold = spec.get("bold", False) if bold is None else bold
    actual_color = color or spec.get("color")
    if actual_color:
        run.font.color.rgb = RGBColor.from_string(actual_color)


def configure_style(style, spec: dict, paragraph_spec: dict | None = None) -> None:
    set_run_font(_style_run_proxy(style), spec)
    if paragraph_spec:
        fmt = style.paragraph_format
        fmt.space_before = Pt(paragraph_spec.get("space_before_pt", 0))
        fmt.space_after = Pt(paragraph_spec.get("space_after_pt", 0))
        if "line_spacing_pt" in paragraph_spec:
            fmt.line_spacing = Pt(paragraph_spec["line_spacing_pt"])
        else:
            fmt.line_spacing = paragraph_spec.get("line_spacing", 1.0)
        fmt.alignment = ALIGNMENTS.get(paragraph_spec.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)


class _style_run_proxy:
    """Small adapter exposing a run-like API for a Word style's rPr."""

    def __init__(self, style):
        self._element = style.element
        self.font = style.font


def configure_document(doc: Document, profile: dict) -> None:
    page = profile["page"]
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Mm(page["width_mm"])
    section.page_height = Mm(page["height_mm"])
    section.top_margin = Mm(page["top_mm"])
    section.bottom_margin = Mm(page["bottom_mm"])
    section.left_margin = Mm(page["left_mm"])
    section.right_margin = Mm(page["right_mm"])
    section.header_distance = Mm(page["header_mm"])
    section.footer_distance = Mm(page["footer_mm"])

    normal = doc.styles["Normal"]
    configure_style(normal, profile["body"], profile["body"])
    normal.paragraph_format.first_line_indent = Pt(
        profile["body"]["size_pt"] * profile["body"].get("first_line_chars", 0)
    )

    title = doc.styles["Title"]
    configure_style(title, profile["title"], profile["title"])
    title.paragraph_format.keep_with_next = True

    for level in range(1, 5):
        style = doc.styles[f"Heading {level}"]
        heading_spec = {**profile["headings"][level], "ascii_font": profile["body"].get("ascii_font", "Calibri")}
        paragraph_spec = {
            "alignment": "left",
            "space_before_pt": 8 if level == 1 else 4,
            "space_after_pt": 2,
            "line_spacing_pt": profile["body"].get("line_spacing_pt", profile["body"]["size_pt"] * 1.5),
        }
        configure_style(style, heading_spec, paragraph_spec)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.keep_with_next = True

    if "Office Metadata" not in [s.name for s in doc.styles]:
        metadata_style = doc.styles.add_style("Office Metadata", WD_STYLE_TYPE.PARAGRAPH)
    else:
        metadata_style = doc.styles["Office Metadata"]
    configure_style(metadata_style, profile["metadata"], {"alignment": "left", "space_after_pt": 3, "line_spacing": 1.0})


def add_field(paragraph, field_name: str, font_spec: dict) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_name} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, font_spec)


def configure_footer(doc: Document, profile: dict, profile_name: str) -> None:
    spec = profile["footer"]
    if spec["style"] == "none":
        return
    section = doc.sections[0]
    if spec["style"] == "official":
        doc.settings.odd_and_even_pages_header_footer = True
        for footer, alignment in (
            (section.footer, WD_ALIGN_PARAGRAPH.RIGHT),
            (section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT),
        ):
            p = footer.paragraphs[0]
            p.alignment = alignment
            prefix = p.add_run("— ")
            set_run_font(prefix, spec)
            add_field(p, "PAGE", spec)
            suffix = p.add_run(" —")
            set_run_font(suffix, spec)
    else:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field(p, "PAGE", spec)


def add_title(doc: Document, title: str, profile: dict) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = ALIGNMENTS[profile["title"].get("alignment", "center")]
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(title)
    set_run_font(run, profile["title"])


def add_recipient(doc: Document, recipient: str, profile: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    text = recipient.rstrip("：:") + "："
    run = p.add_run(text)
    set_run_font(run, profile["recipient"])


def add_body(doc: Document, text: str, profile: dict, *, indent: bool = True, alignment: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = ALIGNMENTS.get(alignment or profile["body"].get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
    indent_chars = profile["body"].get("first_line_chars", 0) if indent else 0
    p.paragraph_format.first_line_indent = Pt(profile["body"]["size_pt"] * indent_chars)
    run = p.add_run(text)
    set_run_font(run, profile["body"])


def add_heading(doc: Document, text: str, profile: dict, level: int = 1) -> None:
    level = min(max(level, 1), 4)
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    spec = {**profile["headings"][level], "ascii_font": profile["body"].get("ascii_font", "Calibri")}
    set_run_font(run, spec)


def add_bullets(doc: Document, items: Iterable[str], profile: dict) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Pt(profile["body"]["size_pt"] * 2)
        p.paragraph_format.first_line_indent = Pt(-profile["body"]["size_pt"])
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(str(item))
        set_run_font(run, profile["body"])


def add_definition_list(doc: Document, values: dict, profile: dict) -> None:
    for label, value in values.items():
        if value in (None, "", []):
            continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        label_run = p.add_run(f"{label}：")
        set_run_font(label_run, profile["body"], bold=True)
        value_run = p.add_run(str(value))
        set_run_font(value_run, profile["body"])


def add_sections(doc: Document, sections: Iterable[dict], profile: dict, *, auto_number: bool = False) -> None:
    for index, section in enumerate(sections, start=1):
        heading = str(section.get("heading", "")).strip()
        if heading:
            if auto_number and not re.match(r"^[一二三四五六七八九十]+、", heading):
                numeral = CN_NUMERALS[index - 1] if index <= len(CN_NUMERALS) else str(index)
                heading = f"{numeral}、{heading}"
            add_heading(doc, heading, profile, int(section.get("level", 1)))
        for paragraph in as_list(section.get("paragraphs") or section.get("body")):
            add_body(doc, paragraph, profile)
        add_bullets(doc, as_list(section.get("bullets")), profile)


def add_attachments(doc: Document, attachments: Iterable[str], profile: dict) -> None:
    attachments = list(attachments)
    if not attachments:
        return
    add_body(doc, f"附件：1. {attachments[0]}", profile, indent=False)
    for index, attachment in enumerate(attachments[1:], start=2):
        add_body(doc, f"　　　{index}. {attachment}", profile, indent=False)


def add_signature(doc: Document, payload: dict, profile: dict) -> None:
    values = [payload.get("department"), payload.get("signer"), payload.get("contact"), payload.get("date")]
    seen = set()
    for value in values:
        if not value or value in seen:
            continue
        seen.add(value)
        p = doc.add_paragraph()
        p.alignment = ALIGNMENTS[profile["signature"].get("alignment", "right")]
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(value))
        set_run_font(run, profile["signature"])


def build_leave_note(doc: Document, payload: dict, profile: dict) -> None:
    add_title(doc, payload.get("title", "请假条"), profile)
    add_recipient(doc, require_or_placeholder(payload, "recipient", "称谓"), profile)
    body = as_list(payload.get("body"))
    if not body:
        reason = require_or_placeholder(payload, "reason", "请假原因")
        start = require_or_placeholder(payload, "start_date", "开始时间")
        end = require_or_placeholder(payload, "end_date", "结束时间")
        duration = require_or_placeholder(payload, "duration", "请假时长")
        leave_type = payload.get("leave_type")
        first = f"本人因{reason}，拟于{start}至{end}请假，共{duration}"
        if leave_type:
            first = f"本人因{reason}，拟申请{leave_type}，时间为{start}至{end}，共{duration}"
        first += "。"
        if payload.get("handover"):
            first += f"请假期间，{payload['handover']}。"
        if payload.get("contact"):
            first += f"如有紧急事项，可通过{payload['contact']}联系本人。"
        body = [first]
    for paragraph in body:
        add_body(doc, paragraph, profile)
    add_body(doc, payload.get("closing", "恳请批准。"), profile)
    if payload.get("complimentary_close"):
        for line in as_list(payload["complimentary_close"]):
            add_body(doc, line, profile, indent=False)
    add_attachments(doc, as_list(payload.get("attachments")), profile)
    add_signature(doc, payload, profile)


def build_request(doc: Document, payload: dict, profile: dict) -> None:
    add_title(doc, require_or_placeholder(payload, "title", "请示标题"), profile)
    add_recipient(doc, require_or_placeholder(payload, "recipient", "主送对象"), profile)
    paragraphs = as_list(payload.get("body"))
    if not paragraphs:
        paragraphs = as_list(payload.get("background")) + as_list(payload.get("proposal")) + as_list(payload.get("request"))
    for paragraph in paragraphs:
        add_body(doc, paragraph, profile)
    add_sections(doc, payload.get("sections", []), profile, auto_number=True)
    add_body(doc, payload.get("closing", "妥否，请批示。"), profile)
    add_attachments(doc, as_list(payload.get("attachments")), profile)
    add_signature(doc, payload, profile)


def build_notice(doc: Document, payload: dict, profile: dict) -> None:
    add_title(doc, require_or_placeholder(payload, "title", "通知标题"), profile)
    add_recipient(doc, require_or_placeholder(payload, "recipient", "主送对象"), profile)
    for paragraph in as_list(payload.get("intro") or payload.get("body")):
        add_body(doc, paragraph, profile)
    add_sections(doc, payload.get("sections", []), profile, auto_number=True)
    add_body(doc, payload.get("closing", "特此通知。"), profile)
    add_attachments(doc, as_list(payload.get("attachments")), profile)
    add_signature(doc, payload, profile)


def build_application(doc: Document, payload: dict, profile: dict) -> None:
    add_title(doc, require_or_placeholder(payload, "title", "申请书标题"), profile)
    add_recipient(doc, require_or_placeholder(payload, "recipient", "受理对象"), profile)
    paragraphs = as_list(payload.get("body"))
    if not paragraphs:
        identity = require_or_placeholder(payload, "identity", "身份或基本情况")
        request = require_or_placeholder(payload, "request", "申请事项")
        reason = require_or_placeholder(payload, "reason", "申请理由")
        paragraphs = [f"{identity}，现申请{request}。申请理由如下：{reason}。"]
        paragraphs += as_list(payload.get("details"))
        if payload.get("commitment"):
            paragraphs.append(str(payload["commitment"]))
    for paragraph in paragraphs:
        add_body(doc, paragraph, profile)
    add_body(doc, payload.get("closing", "恳请审核批准。"), profile)
    add_attachments(doc, as_list(payload.get("attachments")), profile)
    add_signature(doc, payload, profile)


def build_email(doc: Document, payload: dict, profile: dict) -> None:
    subject = require_or_placeholder(payload, "subject", "邮件主题")
    add_title(doc, f"主题：{subject}", profile)
    metadata = {"收件人": require_or_placeholder(payload, "to", "收件人")}
    if payload.get("cc"):
        metadata["抄送"] = payload["cc"]
    for label, value in metadata.items():
        p = doc.add_paragraph(style="Office Metadata")
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(f"{label}：{value}")
        set_run_font(run, profile["metadata"])
    greeting = payload.get("greeting") or require_or_placeholder(payload, "recipient", "称呼")
    add_recipient(doc, greeting, profile)
    for paragraph in as_list(payload.get("paragraphs") or payload.get("body")):
        add_body(doc, paragraph, profile, indent=False, alignment="left")
    if payload.get("action"):
        add_body(doc, str(payload["action"]), profile, indent=False, alignment="left")
    if payload.get("closing"):
        add_body(doc, str(payload["closing"]), profile, indent=False, alignment="left")
    add_attachments(doc, as_list(payload.get("attachments")), profile)
    add_signature(doc, payload, profile)


def add_report_section(doc: Document, heading: str, content: Any, profile: dict, index: int) -> int:
    items = as_list(content)
    if not items:
        return index
    numeral = CN_NUMERALS[index - 1] if index <= len(CN_NUMERALS) else str(index)
    add_heading(doc, f"{numeral}、{heading}", profile, 1)
    if len(items) == 1:
        add_body(doc, items[0], profile)
    else:
        add_bullets(doc, items, profile)
    return index + 1


def build_travel_report(doc: Document, payload: dict, profile: dict) -> None:
    add_title(doc, require_or_placeholder(payload, "title", "出差报告标题"), profile)
    index = 1
    if payload.get("summary"):
        index = add_report_section(doc, "核心结论", payload["summary"], profile, index)
    if payload.get("basics"):
        numeral = CN_NUMERALS[index - 1]
        add_heading(doc, f"{numeral}、基本信息", profile, 1)
        add_definition_list(doc, payload["basics"], profile)
        index += 1
    for heading, key in (
        ("行程与工作开展", "itinerary"),
        ("主要成果与收获", "results"),
        ("问题与风险", "issues"),
        ("建议与后续安排", "suggestions"),
        ("需跟踪落实事项", "follow_up"),
        ("费用说明", "expenses"),
    ):
        index = add_report_section(doc, heading, payload.get(key), profile, index)
    add_sections(doc, payload.get("sections", []), profile, auto_number=True)
    add_attachments(doc, as_list(payload.get("attachments")), profile)
    add_signature(doc, payload, profile)


def build_work_report(doc: Document, payload: dict, profile: dict) -> None:
    add_title(doc, require_or_placeholder(payload, "title", "报告标题"), profile)
    index = 1
    for heading, key in (
        ("本期核心结论", "summary"),
        ("重点工作完成情况", "achievements"),
        ("问题与风险", "issues"),
        ("下一阶段计划", "plans"),
        ("需协调支持事项", "support"),
    ):
        index = add_report_section(doc, heading, payload.get(key), profile, index)
    add_sections(doc, payload.get("sections", []), profile, auto_number=True)
    add_attachments(doc, as_list(payload.get("attachments")), profile)
    add_signature(doc, payload, profile)


def build_official_document(doc: Document, payload: dict, profile: dict) -> None:
    add_title(doc, require_or_placeholder(payload, "title", "公文标题"), profile)
    if payload.get("recipient"):
        add_recipient(doc, str(payload["recipient"]), profile)
    for paragraph in as_list(payload.get("intro") or payload.get("body")):
        add_body(doc, paragraph, profile)
    add_sections(doc, payload.get("sections", []), profile, auto_number=True)
    if payload.get("closing"):
        add_body(doc, str(payload["closing"]), profile)
    add_attachments(doc, as_list(payload.get("attachments")), profile)
    add_signature(doc, payload, profile)


BUILDERS = {
    "leave-note": build_leave_note,
    "request": build_request,
    "notice": build_notice,
    "application": build_application,
    "formal-email": build_email,
    "travel-report": build_travel_report,
    "weekly-report": build_work_report,
    "monthly-report": build_work_report,
    "work-report": build_work_report,
    "official-document": build_official_document,
}


def load_payload(path: str) -> dict:
    if path == "-":
        return json.load(sys.stdin)
    with Path(path).open("r", encoding="utf-8") as handle:
        payload = json.load(handle)
    if not isinstance(payload, dict):
        raise ValueError("Input JSON must be an object")
    return payload


def generate(document_type: str, payload: dict, profile_name: str, output: Path) -> dict:
    resolved_name, profile = resolve_profile(document_type, profile_name, payload)
    doc = Document()
    configure_document(doc, profile)
    configure_footer(doc, profile, resolved_name)
    doc.core_properties.title = str(payload.get("title") or payload.get("subject") or document_type)
    doc.core_properties.subject = f"Generated with office-document-templates ({document_type})"
    doc.core_properties.author = str(payload.get("signer") or "")
    BUILDERS[document_type](doc, payload, profile)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    warnings = []
    serialized = json.dumps(payload, ensure_ascii=False)
    if "[待补充" in serialized:
        warnings.append("Payload contains [待补充] placeholders; confirm before final submission.")
    return {
        "output": str(output.resolve()),
        "document_type": document_type,
        "profile": resolved_name,
        "warnings": warnings,
        "format_source_keys": list(FORMAT_SOURCES),
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--type", choices=DOCUMENT_TYPES)
    parser.add_argument("--input", help="UTF-8 JSON payload path, or - for stdin")
    parser.add_argument("--output", help="Destination .docx path")
    parser.add_argument("--profile", default="auto", choices=("auto", *PROFILES))
    parser.add_argument("--list-types", action="store_true")
    parser.add_argument("--list-profiles", action="store_true")
    parser.add_argument("--print-schema", choices=DOCUMENT_TYPES)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if args.list_types:
        print("\n".join(DOCUMENT_TYPES))
        return 0
    if args.list_profiles:
        print(json.dumps({name: PROFILES[name] for name in PROFILES}, ensure_ascii=False, indent=2))
        return 0
    if args.print_schema:
        print(json.dumps(SCHEMAS[args.print_schema], ensure_ascii=False, indent=2))
        return 0
    if not args.type or not args.input or not args.output:
        raise SystemExit("--type, --input, and --output are required for generation")
    output = Path(args.output)
    if output.suffix.lower() != ".docx":
        raise SystemExit("--output must end with .docx")
    result = generate(args.type, load_payload(args.input), args.profile, output)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
